"""Tests for Cochrane-style meta-analysis DOCX generation.

Following TDD: Tests written before implementation.
Danish labels: "Inkluderede studier", "Risiko for bias", "Evidens-syntese"
"""
from __future__ import annotations

import pytest


class TestMetaAnalysisDocxWriterExists:
    """Tests verifying meta-analysis DOCX writer function exists."""

    def test_write_meta_analysis_docx_importable(self) -> None:
        """write_meta_analysis_docx should be importable."""
        from procedurewriter.pipeline.docx_writer import write_meta_analysis_docx

        assert write_meta_analysis_docx is not None

    def test_function_accepts_synthesis_output(self) -> None:
        """Function should accept SynthesisOutput as parameter."""
        # Should not raise
        import inspect

        from procedurewriter.pipeline.docx_writer import write_meta_analysis_docx
        sig = inspect.signature(write_meta_analysis_docx)
        params = list(sig.parameters.keys())
        assert "synthesis" in params or "output" in params


class TestCochraneSections:
    """Tests for Cochrane-standard sections in generated DOCX."""

    @pytest.fixture
    def synthesis_output(self):
        """Create sample synthesis output for testing."""
        from procedurewriter.agents.meta_analysis.synthesizer_agent import (
            ForestPlotEntry,
            HeterogeneityMetrics,
            PooledEstimate,
            SynthesisOutput,
        )

        return SynthesisOutput(
            pooled_estimate=PooledEstimate(
                pooled_effect=0.65,
                ci_lower=0.45,
                ci_upper=0.85,
                effect_size_type="OR",
                p_value=0.003,
                se=0.1,
            ),
            heterogeneity=HeterogeneityMetrics(
                cochrans_q=8.5,
                i_squared=52.3,
                tau_squared=0.05,
                df=3,
                p_value=0.037,
                interpretation="moderate",
            ),
            included_studies=4,
            total_sample_size=2400,
            grade_summary="Moderate certainty evidence suggests benefit.",
            forest_plot_data=[
                ForestPlotEntry(
                    study_id="S1",
                    study_label="Jensen 2022",
                    effect_size=0.72,
                    ci_lower=0.48,
                    ci_upper=0.96,
                    weight=0.28,
                    sample_size=600,
                ),
                ForestPlotEntry(
                    study_id="S2",
                    study_label="Hansen 2023",
                    effect_size=0.58,
                    ci_lower=0.38,
                    ci_upper=0.78,
                    weight=0.35,
                    sample_size=800,
                ),
            ],
        )

    @pytest.fixture
    def orchestrator_output(self, synthesis_output):
        """Create sample orchestrator output for testing."""
        from procedurewriter.agents.meta_analysis.orchestrator import OrchestratorOutput

        return OrchestratorOutput(
            synthesis=synthesis_output,
            included_study_ids=["S1", "S2", "S3", "S4"],
            excluded_study_ids=["S5"],
            exclusion_reasons={"S5": "Population mismatch"},
            manual_review_needed=["S6"],
        )

    def test_docx_has_danish_section_headers(self, tmp_path, orchestrator_output) -> None:
        """Generated DOCX should have Danish section headers."""
        from docx import Document

        from procedurewriter.pipeline.docx_writer import write_meta_analysis_docx

        output_path = tmp_path / "meta_analysis.docx"

        write_meta_analysis_docx(
            output=orchestrator_output,
            output_path=output_path,
            run_id="test-run-123",
        )

        doc = Document(str(output_path))
        text = "\n".join([p.text for p in doc.paragraphs])

        # Danish labels as specified
        assert "Inkluderede studier" in text
        assert "Risiko for bias" in text
        assert "Evidens-syntese" in text

    def test_docx_contains_pooled_estimate(self, tmp_path, orchestrator_output) -> None:
        """DOCX should contain pooled effect estimate."""
        from docx import Document

        from procedurewriter.pipeline.docx_writer import write_meta_analysis_docx

        output_path = tmp_path / "meta_analysis.docx"

        write_meta_analysis_docx(
            output=orchestrator_output,
            output_path=output_path,
            run_id="test-run-123",
        )

        doc = Document(str(output_path))
        # Get text from paragraphs and tables
        text_parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        text = "\n".join(text_parts)

        # Should contain pooled estimate value
        assert "0.65" in text or "0,65" in text  # Danish may use comma

    def test_docx_contains_heterogeneity_metrics(self, tmp_path, orchestrator_output) -> None:
        """DOCX should contain I² heterogeneity statistic."""
        from docx import Document

        from procedurewriter.pipeline.docx_writer import write_meta_analysis_docx

        output_path = tmp_path / "meta_analysis.docx"

        write_meta_analysis_docx(
            output=orchestrator_output,
            output_path=output_path,
            run_id="test-run-123",
        )

        doc = Document(str(output_path))
        # Get text from paragraphs and tables
        text_parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        text = "\n".join(text_parts)

        # Should contain I² value
        assert "I²" in text or "I2" in text
        assert "52" in text  # I² = 52.3%

    def test_docx_contains_grade_summary(self, tmp_path, orchestrator_output) -> None:
        """DOCX should contain GRADE certainty summary."""
        from docx import Document

        from procedurewriter.pipeline.docx_writer import write_meta_analysis_docx

        output_path = tmp_path / "meta_analysis.docx"

        write_meta_analysis_docx(
            output=orchestrator_output,
            output_path=output_path,
            run_id="test-run-123",
        )

        doc = Document(str(output_path))
        text = "\n".join([p.text for p in doc.paragraphs])

        # Should contain GRADE summary
        assert "Moderate" in text or "certainty" in text.lower()

    def test_docx_lists_excluded_studies(self, tmp_path, orchestrator_output) -> None:
        """DOCX should list excluded studies with reasons."""
        from docx import Document

        from procedurewriter.pipeline.docx_writer import write_meta_analysis_docx

        output_path = tmp_path / "meta_analysis.docx"

        write_meta_analysis_docx(
            output=orchestrator_output,
            output_path=output_path,
            run_id="test-run-123",
        )

        doc = Document(str(output_path))
        text = "\n".join([p.text for p in doc.paragraphs])

        # Should list excluded study
        assert "S5" in text
        assert "Population mismatch" in text


class TestForestPlotTable:
    """Tests for forest plot data in DOCX table format."""

    @pytest.fixture
    def synthesis_output(self):
        """Create sample synthesis output with forest plot data."""
        from procedurewriter.agents.meta_analysis.synthesizer_agent import (
            ForestPlotEntry,
            HeterogeneityMetrics,
            PooledEstimate,
            SynthesisOutput,
        )

        return SynthesisOutput(
            pooled_estimate=PooledEstimate(
                pooled_effect=0.65,
                ci_lower=0.45,
                ci_upper=0.85,
                effect_size_type="OR",
                p_value=0.003,
                se=0.1,
            ),
            heterogeneity=HeterogeneityMetrics(
                cochrans_q=8.5,
                i_squared=52.3,
                tau_squared=0.05,
                df=3,
                p_value=0.037,
                interpretation="moderate",
            ),
            included_studies=2,
            total_sample_size=1400,
            grade_summary="Moderate certainty.",
            forest_plot_data=[
                ForestPlotEntry(
                    study_id="S1",
                    study_label="Jensen 2022",
                    effect_size=0.72,
                    ci_lower=0.48,
                    ci_upper=0.96,
                    weight=0.45,
                    sample_size=600,
                ),
                ForestPlotEntry(
                    study_id="S2",
                    study_label="Hansen 2023",
                    effect_size=0.58,
                    ci_lower=0.38,
                    ci_upper=0.78,
                    weight=0.55,
                    sample_size=800,
                ),
            ],
        )

    def test_docx_has_forest_plot_table(self, tmp_path, synthesis_output) -> None:
        """DOCX should contain a table with forest plot data."""
        from docx import Document

        from procedurewriter.agents.meta_analysis.orchestrator import OrchestratorOutput
        from procedurewriter.pipeline.docx_writer import write_meta_analysis_docx

        output_path = tmp_path / "meta_analysis.docx"
        orchestrator_output = OrchestratorOutput(
            synthesis=synthesis_output,
            included_study_ids=["S1", "S2"],
            excluded_study_ids=[],
            exclusion_reasons={},
            manual_review_needed=[],
        )

        write_meta_analysis_docx(
            output=orchestrator_output,
            output_path=output_path,
            run_id="test-run-123",
        )

        doc = Document(str(output_path))

        # Should have at least one table
        assert len(doc.tables) >= 1

    def test_forest_table_has_study_rows(self, tmp_path, synthesis_output) -> None:
        """Forest plot table should have rows for each study."""
        from docx import Document

        from procedurewriter.agents.meta_analysis.orchestrator import OrchestratorOutput
        from procedurewriter.pipeline.docx_writer import write_meta_analysis_docx

        output_path = tmp_path / "meta_analysis.docx"
        orchestrator_output = OrchestratorOutput(
            synthesis=synthesis_output,
            included_study_ids=["S1", "S2"],
            excluded_study_ids=[],
            exclusion_reasons={},
            manual_review_needed=[],
        )

        write_meta_analysis_docx(
            output=orchestrator_output,
            output_path=output_path,
            run_id="test-run-123",
        )

        doc = Document(str(output_path))
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text for cell in row.cells)
                table_texts.append(row_text)

        all_table_text = "\n".join(table_texts)

        # Should contain study labels
        assert "Jensen 2022" in all_table_text
        assert "Hansen 2023" in all_table_text


class TestPRISMAFlowchart:
    """Tests for PRISMA 2020 flowchart data."""

    def test_prisma_data_included_in_output(self) -> None:
        """PRISMA flowchart counts should be calculable from output."""
        from procedurewriter.agents.meta_analysis.orchestrator import OrchestratorOutput
        from procedurewriter.agents.meta_analysis.synthesizer_agent import (
            HeterogeneityMetrics,
            PooledEstimate,
            SynthesisOutput,
        )

        synthesis = SynthesisOutput(
            pooled_estimate=PooledEstimate(
                pooled_effect=0.5,
                ci_lower=0.3,
                ci_upper=0.7,
                effect_size_type="OR",
                p_value=0.01,
            ),
            heterogeneity=HeterogeneityMetrics(
                cochrans_q=5.0,
                i_squared=40.0,
                tau_squared=0.02,
                df=2,
                p_value=0.1,
            ),
            included_studies=3,
            total_sample_size=1500,
            grade_summary="Test",
            forest_plot_data=[],
        )

        output = OrchestratorOutput(
            synthesis=synthesis,
            included_study_ids=["S1", "S2", "S3"],
            excluded_study_ids=["S4", "S5"],
            exclusion_reasons={"S4": "Wrong population", "S5": "Wrong outcome"},
            manual_review_needed=["S6"],
        )

        # PRISMA counts should be derivable
        identified = len(output.included_study_ids) + len(output.excluded_study_ids) + len(output.manual_review_needed)
        screened = identified
        excluded = len(output.excluded_study_ids)
        included = len(output.included_study_ids)

        assert identified == 6
        assert excluded == 2
        assert included == 3
