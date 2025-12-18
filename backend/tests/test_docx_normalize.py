from __future__ import annotations

from pathlib import Path

from docx import Document

from procedurewriter.pipeline.normalize import extract_docx_blocks, normalize_docx_blocks


def test_extract_and_normalize_docx_blocks(tmp_path: Path) -> None:
    doc = Document()
    doc.add_heading("My Title", level=1)
    doc.add_paragraph("Intro paragraph.")
    doc.add_paragraph("First bullet", style="List Bullet")
    doc.add_paragraph("Second bullet", style="List Bullet")
    doc.add_paragraph("Step one", style="List Number")

    path = tmp_path / "sample.docx"
    doc.save(str(path))

    blocks = extract_docx_blocks(path)
    assert any(b["kind"] == "heading" and "My Title" in b["text"] for b in blocks)
    assert any(b["kind"] == "bullet" and "First bullet" in b["text"] for b in blocks)
    assert any(b["kind"] == "numbered" and "Step one" in b["text"] for b in blocks)

    normalized = normalize_docx_blocks(blocks)
    assert "# My Title" in normalized
    assert "- First bullet" in normalized
    assert "1. Step one" in normalized
