from pathlib import Path

from docx import Document

from procedurewriter.pipeline.docx_writer import write_procedure_docx
from procedurewriter.pipeline.types import SourceRecord


def test_docx_writer_smoke(tmp_path: Path):
    md = (
        "# Procedure: Test\n\n"
        "## Indikation\n"
        "Sætning. [S:SRC0001]\n\n"
        "## Sikkerhedsboks\n"
        "- OBS: Stop hvis noget. [S:SRC0001]\n\n"
        "## Fremgangsmåde\n"
        "1. Trin et. [S:SRC0001]\n"
    )
    sources = [
        SourceRecord(
            source_id="SRC0001",
            fetched_at_utc="2020-01-01T00:00:00+00:00",
            kind="dummy",
            title="Dummy",
            year=2020,
            url=None,
            doi=None,
            pmid=None,
            raw_path=str(tmp_path / "raw.txt"),
            normalized_path=str(tmp_path / "norm.txt"),
            raw_sha256="x",
            normalized_sha256="y",
            extraction_notes=None,
            terms_licence_note=None,
            extra={},
        )
    ]
    out = tmp_path / "Procedure.docx"
    write_procedure_docx(markdown_text=md, sources=sources, output_path=out, run_id="RID", manifest_hash="H")
    assert out.exists()

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Procedure: Test" in text
    assert "Referencer" in text
    assert any(p.style.name == "List Number" for p in doc.paragraphs)
    assert len(doc.tables) >= 1
    assert any(
        (run.text == "[S:SRC0001]" and bool(run.font.superscript))
        for p in doc.paragraphs
        for run in p.runs
    )
