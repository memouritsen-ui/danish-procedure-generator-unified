"""DOCX Writer with template-based customization.

Generates procedure documents with configurable structure, styling, and content options.
"""
from __future__ import annotations

import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from procedurewriter.config_store import load_yaml
from procedurewriter.pipeline.types import SourceRecord

_citation_tag_re = re.compile(r"\[S:[^\]]+\]")


def format_grade_badge(certainty_level: str) -> str:
    """Format GRADE certainty as visual badge for Markdown/DOCX.

    Args:
        certainty_level: One of "High", "Moderate", "Low", "Very Low".

    Returns:
        Formatted badge string: **[GRADE: Level]**
    """
    return f"**[GRADE: {certainty_level}]**"


def load_docx_template(template_path: Path | None = None) -> dict[str, Any]:
    """Load DOCX template configuration.

    Args:
        template_path: Path to docx_template.yaml. If None, uses defaults.

    Returns:
        Template configuration dictionary
    """
    if template_path and template_path.exists():
        config = load_yaml(template_path)
        if isinstance(config, dict):
            return config

    # Return default configuration
    return {
        "structure": {
            "sections": [
                {"id": "indikation", "heading": "Indikation", "visible": True, "format": "bullets"},
                {"id": "kontraindikation", "heading": "Kontraindikation", "visible": True, "format": "bullets"},
                {"id": "udstyr", "heading": "Udstyr", "visible": True, "format": "bullets"},
                {"id": "forberedelse", "heading": "Forberedelse", "visible": True, "format": "numbered"},
                {"id": "procedure", "heading": "Procedure", "visible": True, "format": "numbered"},
                {"id": "sikkerhedsboks", "heading": "Sikkerhedsboks", "visible": True, "format": "safety_box"},
                {"id": "komplikationer", "heading": "Komplikationer", "visible": True, "format": "bullets"},
                {"id": "efterbehandling", "heading": "Efterbehandling", "visible": True, "format": "bullets"},
                {"id": "dokumentation", "heading": "Dokumentation", "visible": True, "format": "bullets"},
                {"id": "referencer", "heading": "Referencer", "visible": True, "format": "references"},
            ]
        },
        "styling": {
            "fonts": {
                "body": {"family": "Calibri", "size": 11},
                "heading1": {"family": "Calibri", "size": 16, "bold": True},
                "heading2": {"family": "Calibri", "size": 14, "bold": True},
            },
            "colors": {
                "heading1": "#003366",
                "heading2": "#003366",
                "body": "#000000",
                "citation": "#6e6e6e",
                "safety_box_background": "#FFF2CC",
            },
        },
        "content": {
            "citations": {"style": "superscript", "size": 8, "show_references": True},
            "evidence_badges": {"show_in_text": False, "show_in_references": True},
            "audit_trail": {"show": True, "abbreviated_hash": True},
            "page_numbers": {"show": True, "format": "Side X af Y"},
        },
    }


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def _apply_font_style(run: Any, config: dict[str, Any], style_name: str = "body") -> None:
    """Apply font styling from template config."""
    fonts = config.get("styling", {}).get("fonts", {})
    font_config = fonts.get(style_name, fonts.get("body", {}))

    if "family" in font_config:
        run.font.name = font_config["family"]
    if "size" in font_config:
        run.font.size = Pt(font_config["size"])
    if font_config.get("bold"):
        run.font.bold = True
    if font_config.get("italic"):
        run.font.italic = True


def _apply_heading_color(paragraph: Any, config: dict[str, Any], level: int) -> None:
    """Apply heading color from template config."""
    colors = config.get("styling", {}).get("colors", {})
    color_key = f"heading{level}"
    hex_color = colors.get(color_key, "#003366")

    for run in paragraph.runs:
        r, g, b = _hex_to_rgb(hex_color)
        run.font.color.rgb = RGBColor(r, g, b)


def write_procedure_docx(
    *,
    markdown_text: str,
    sources: list[SourceRecord],
    output_path: Path,
    run_id: str,
    manifest_hash: str,
    template_path: Path | None = None,
    quality_score: int | None = None,
) -> None:
    """Write procedure document with template-based customization.

    Args:
        markdown_text: Procedure content in markdown format
        sources: List of source records for references
        output_path: Path to save the DOCX file
        run_id: Unique run identifier
        manifest_hash: SHA256 hash of the manifest
        template_path: Optional path to docx_template.yaml
        quality_score: Optional quality score (1-10)
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Load template configuration
    config = load_docx_template(template_path)

    doc: Any = Document()

    # Apply document margins
    margins = config.get("styling", {}).get("margins", {})
    for section in doc.sections:
        section.top_margin = Inches(margins.get("top", 1.0))
        section.bottom_margin = Inches(margins.get("bottom", 1.0))
        section.left_margin = Inches(margins.get("left", 1.0))
        section.right_margin = Inches(margins.get("right", 1.0))

    # Add branding/logo if configured
    branding = config.get("styling", {}).get("branding", {})
    logo_placement = branding.get("logo_placement", "none")
    logo_path = branding.get("logo_path", "")

    if logo_placement == "header" and logo_path and Path(logo_path).exists():
        header = doc.sections[0].header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(branding.get("logo_width", 1.5)))

    # Render markdown content
    _render_markdown_with_template(doc, markdown_text, config)

    # Add references section if configured
    content_config = config.get("content", {})
    citations_config = content_config.get("citations", {})
    if citations_config.get("show_references", True):
        doc.add_page_break()
        ref_heading = doc.add_heading("Referencer", level=1)
        _apply_heading_color(ref_heading, config, 1)

        evidence_badges = content_config.get("evidence_badges", {})
        show_badges = evidence_badges.get("show_in_references", True)

        for src in sources:
            parts: list[str] = []
            if src.title:
                parts.append(src.title)
            if src.year:
                parts.append(str(src.year))
            if src.url:
                parts.append(src.url)
            if src.doi:
                parts.append(f"DOI: {src.doi}")
            if src.pmid:
                parts.append(f"PMID: {src.pmid}")

            p = doc.add_paragraph()
            rid = p.add_run(f"[{src.source_id}] ")
            rid.bold = True
            p.add_run(" — ".join(parts))

            # Add evidence badge if configured
            if show_badges and src.extra:
                badge = src.extra.get("evidence_badge")
                badge_color = src.extra.get("evidence_badge_color")
                if badge:
                    badge_run = p.add_run(f" [{badge}]")
                    badge_run.font.size = Pt(8)
                    if badge_color:
                        r, g, b = _hex_to_rgb(badge_color)
                        badge_run.font.color.rgb = RGBColor(r, g, b)

    # Add footer with audit trail
    audit_config = content_config.get("audit_trail", {})
    if audit_config.get("show", True):
        ts = datetime.now(UTC).replace(microsecond=0).isoformat()
        hash_display = manifest_hash[:8] if audit_config.get("abbreviated_hash", True) else manifest_hash

        footer_parts = [f"run_id={run_id}", f"timestamp_utc={ts}", f"manifest_sha256={hash_display}"]

        # Add quality score if configured
        quality_config = content_config.get("quality_score", {})
        if quality_config.get("show", False) and quality_score is not None:
            footer_parts.append(f"quality_score={quality_score}/10")

        footer_text = " | ".join(footer_parts)

        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = footer_text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page numbers if configured
    page_config = content_config.get("page_numbers", {})
    if page_config.get("show", True):
        _add_page_numbers(doc, page_config.get("format", "Side X af Y"))

    doc.save(str(output_path))


def _render_markdown_with_template(doc: Any, markdown_text: str, config: dict[str, Any]) -> None:
    """Render markdown content with template-based formatting."""
    current_h2: str | None = None
    safety_lines: list[str] = []

    # Get section visibility settings
    structure = config.get("structure", {})
    sections = structure.get("sections", [])
    section_config = {s["id"]: s for s in sections if isinstance(s, dict)}

    # Get styling colors
    colors = config.get("styling", {}).get("colors", {})
    citation_color = colors.get("citation", "#6e6e6e")
    citation_size = config.get("content", {}).get("citations", {}).get("size", 8)

    def get_section_id(heading: str) -> str | None:
        """Map heading text to section ID."""
        heading_lower = heading.lower().strip()
        for sec in sections:
            if isinstance(sec, dict):
                if sec.get("heading", "").lower() == heading_lower:
                    return sec.get("id")
                if sec.get("id", "").lower() == heading_lower:
                    return sec.get("id")
        return None

    def is_section_visible(section_id: str | None) -> bool:
        """Check if section should be visible."""
        if section_id is None:
            return True
        sec = section_config.get(section_id)
        if sec is None:
            return True
        return sec.get("visible", True)

    def flush_safety_box() -> None:
        nonlocal safety_lines
        if not safety_lines:
            return

        safety_bg = colors.get("safety_box_background", "FFF2CC")
        if safety_bg.startswith("#"):
            safety_bg = safety_bg[1:]

        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        _shade_cell(cell, fill=safety_bg)
        cell.text = ""
        for item in safety_lines:
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.add_run("• ").bold = False
            _add_text_with_citations(p, item, citation_color, citation_size)
        safety_lines = []

    current_section_visible = True

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        if line.startswith("# "):
            flush_safety_box()
            heading_text = line[2:].strip()
            h = doc.add_heading(heading_text, level=0)
            _apply_heading_color(h, config, 1)
            current_h2 = None
            current_section_visible = True

        elif line.startswith("## "):
            flush_safety_box()
            heading_text = line[3:].strip()
            section_id = get_section_id(heading_text)
            current_section_visible = is_section_visible(section_id)

            if current_section_visible:
                current_h2 = heading_text
                h = doc.add_heading(heading_text, level=1)
                _apply_heading_color(h, config, 2)
            else:
                current_h2 = None

        elif line.startswith("### "):
            if not current_section_visible:
                continue
            flush_safety_box()
            heading_text = line[4:].strip()
            h = doc.add_heading(heading_text, level=2)
            _apply_heading_color(h, config, 3)

        elif line.startswith("  - "):
            # Nested bullet (second level)
            if not current_section_visible:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.add_run("○ ").bold = False
            _add_text_with_citations(p, line[4:].strip(), citation_color, citation_size)

        elif line.startswith("- "):
            if not current_section_visible:
                continue
            if _in_safety_box(current_h2):
                safety_lines.append(line[2:].strip())
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.add_run("• ").bold = False
                _add_text_with_citations(p, line[2:].strip(), citation_color, citation_size)

        elif _is_numbered_list_item(line):
            if not current_section_visible:
                continue
            body = line.split(".", 1)[1].strip()
            if _in_safety_box(current_h2):
                safety_lines.append(body)
            else:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                _add_text_with_citations(p, body, citation_color, citation_size)

        else:
            if not current_section_visible:
                continue
            if _in_safety_box(current_h2):
                safety_lines.append(line.strip())
            else:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                _add_text_with_citations(p, line.strip(), citation_color, citation_size)

    flush_safety_box()


def _is_numbered_list_item(line: str) -> bool:
    """Check if line is a numbered list item."""
    if "." not in line:
        return False
    prefix, _rest = line.split(".", 1)
    prefix = prefix.strip()
    return prefix.isdigit()


def _in_safety_box(current_h2: str | None) -> bool:
    """Check if currently in safety box section."""
    return (current_h2 or "").strip().lower() == "sikkerhedsboks"


def _add_text_with_citations(
    paragraph: Any,
    text: str,
    citation_color: str = "#6e6e6e",
    citation_size: int = 8,
) -> None:
    """Add text with styled citation tags."""
    r, g, b = _hex_to_rgb(citation_color)
    pos = 0
    for m in _citation_tag_re.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tag = m.group(0)
        run = paragraph.add_run(tag)
        run.font.superscript = True
        run.font.size = Pt(citation_size)
        run.font.color.rgb = RGBColor(r, g, b)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _shade_cell(cell: Any, *, fill: str) -> None:
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_page_numbers(doc: Any, format_string: str = "Side X af Y") -> None:
    """Add page numbers to document footer."""
    section = doc.sections[0]
    footer = section.footer

    # Get or create paragraph
    if footer.paragraphs:
        p = footer.paragraphs[0]
        if p.text:
            # Add a new paragraph for page numbers
            p = footer.add_paragraph()
    else:
        p = footer.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Parse format string and add page number fields
    if "X af Y" in format_string or "X of Y" in format_string:
        prefix = format_string.split("X")[0]
        suffix = format_string.split("Y")[-1] if "Y" in format_string else ""

        if prefix:
            p.add_run(prefix)

        # Add PAGE field
        _add_field(p, "PAGE")

        middle = " af " if "af" in format_string else " of "
        p.add_run(middle)

        # Add NUMPAGES field
        _add_field(p, "NUMPAGES")

        if suffix:
            p.add_run(suffix)
    else:
        # Simple page number
        prefix = format_string.replace("X", "").strip()
        if prefix:
            p.add_run(prefix + " ")
        _add_field(p, "PAGE")


def _add_field(paragraph: Any, field_code: str) -> None:
    """Add a Word field (PAGE, NUMPAGES, etc.) to paragraph."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.text = field_code
    run._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    # Placeholder text
    run._r.append(OxmlElement("w:t"))

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


# =============================================================================
# Meta-Analysis DOCX Generation (Cochrane-style)
# =============================================================================


def write_meta_analysis_docx(
    *,
    output: Any,  # OrchestratorOutput
    output_path: Path,
    run_id: str,
) -> None:
    """Write Cochrane-style meta-analysis report as DOCX.

    Danish labels used throughout:
    - "Inkluderede studier" (Included studies)
    - "Risiko for bias" (Risk of bias)
    - "Evidens-syntese" (Evidence synthesis)

    Args:
        output: OrchestratorOutput with synthesis results.
        output_path: Path to save the DOCX file.
        run_id: Unique run identifier.
    """
    from datetime import UTC, datetime

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc: Any = Document()

    # Apply document margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    title = doc.add_heading("Meta-analyse Rapport", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with run ID
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"Run ID: {run_id}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Timestamp
    ts = datetime.now(UTC).replace(microsecond=0).isoformat()
    ts_para = doc.add_paragraph()
    ts_run = ts_para.add_run(f"Genereret: {ts}")
    ts_run.font.size = Pt(10)
    ts_run.font.color.rgb = RGBColor(128, 128, 128)
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # Section 1: Inkluderede studier (Included studies)
    doc.add_heading("Inkluderede studier", level=1)

    n_included = len(output.included_study_ids)
    n_excluded = len(output.excluded_study_ids)
    n_manual = len(output.manual_review_needed)

    summary_para = doc.add_paragraph()
    summary_para.add_run(f"Antal inkluderede studier: {n_included}\n")
    summary_para.add_run(f"Antal ekskluderede studier: {n_excluded}\n")
    if n_manual > 0:
        summary_para.add_run(f"Kræver manuel gennemgang: {n_manual}\n")

    # Included study IDs
    if output.included_study_ids:
        doc.add_heading("Inkluderede studie-ID'er", level=2)
        for study_id in output.included_study_ids:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(f"• {study_id}")

    # Excluded studies with reasons
    if output.excluded_study_ids:
        doc.add_heading("Ekskluderede studier", level=2)
        for study_id in output.excluded_study_ids:
            reason = output.exclusion_reasons.get(study_id, "Ingen grund angivet")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(f"• {study_id}: ").bold = True
            p.add_run(reason)

    # Section 2: Risiko for bias (Risk of bias)
    doc.add_page_break()
    doc.add_heading("Risiko for bias", level=1)

    rob_intro = doc.add_paragraph()
    rob_intro.add_run(
        "Vurdering af risiko for bias følger Cochrane Risk of Bias 2.0 værktøjet. "
        "Domæner inkluderer: randomisering, afvigelser fra intervention, manglende data, "
        "måling af udfald og selektion af rapporteret resultat."
    )

    # Section 3: Evidens-syntese (Evidence synthesis)
    doc.add_page_break()
    doc.add_heading("Evidens-syntese", level=1)

    synthesis = output.synthesis

    # Pooled estimate
    doc.add_heading("Samlet effektestimat", level=2)
    pooled = synthesis.pooled_estimate

    pooled_table = doc.add_table(rows=5, cols=2)
    pooled_table.style = "Table Grid"

    pooled_rows = [
        ("Effektmål", f"{pooled.effect_size_type}"),
        ("Samlet effekt", f"{pooled.pooled_effect:.3f}"),
        ("95% CI", f"[{pooled.ci_lower:.3f}, {pooled.ci_upper:.3f}]"),
        ("P-værdi", f"{pooled.p_value:.4f}"),
        ("Antal studier", str(synthesis.included_studies)),
    ]

    for i, (label, value) in enumerate(pooled_rows):
        pooled_table.cell(i, 0).text = label
        pooled_table.cell(i, 1).text = value
        pooled_table.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()  # Spacer

    # Heterogeneity
    doc.add_heading("Heterogenitet", level=2)
    het = synthesis.heterogeneity

    het_table = doc.add_table(rows=5, cols=2)
    het_table.style = "Table Grid"

    het_rows = [
        ("I²", f"{het.i_squared:.1f}%"),
        ("Cochran's Q", f"{het.cochrans_q:.2f}"),
        ("τ²", f"{het.tau_squared:.4f}"),
        ("Frihedsgrader (df)", str(het.df)),
        ("Fortolkning", het.interpretation.capitalize() if het.interpretation else "N/A"),
    ]

    for i, (label, value) in enumerate(het_rows):
        het_table.cell(i, 0).text = label
        het_table.cell(i, 1).text = value
        het_table.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()  # Spacer

    # Forest plot data table
    if synthesis.forest_plot_data:
        doc.add_heading("Forest Plot Data", level=2)

        forest_table = doc.add_table(rows=len(synthesis.forest_plot_data) + 1, cols=5)
        forest_table.style = "Table Grid"

        # Header row
        headers = ["Studie", "Effekt", "95% CI", "Vægt", "N"]
        for j, header in enumerate(headers):
            cell = forest_table.cell(0, j)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        for i, entry in enumerate(synthesis.forest_plot_data, start=1):
            forest_table.cell(i, 0).text = entry.study_label
            forest_table.cell(i, 1).text = f"{entry.effect_size:.3f}"
            forest_table.cell(i, 2).text = f"[{entry.ci_lower:.2f}, {entry.ci_upper:.2f}]"
            forest_table.cell(i, 3).text = f"{entry.weight * 100:.1f}%"
            forest_table.cell(i, 4).text = str(entry.sample_size)

    # GRADE summary
    doc.add_paragraph()  # Spacer
    doc.add_heading("GRADE Vurdering", level=2)

    grade_para = doc.add_paragraph()
    grade_para.add_run(synthesis.grade_summary)

    # Total sample size
    doc.add_paragraph()
    total_n_para = doc.add_paragraph()
    total_n_para.add_run(f"Total stikprøvestørrelse: {synthesis.total_sample_size}")

    # Footer with run ID
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = f"Meta-analyse run_id={run_id}"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(output_path))


# =============================================================================
# Verbose Documentation DOCX Generation
# =============================================================================


def write_source_analysis_docx(
    *,
    sources: list[SourceRecord],
    procedure: str,
    run_id: str,
    output_path: Path,
    search_terms: list[str] | None = None,
) -> None:
    """Write source analysis documentation as DOCX.

    Hybrid format: Readable prose for clinicians + technical appendix for audit.

    Args:
        sources: List of source records from the pipeline.
        procedure: Name of the procedure being documented.
        run_id: Unique run identifier.
        output_path: Path to save the DOCX file.
        search_terms: Optional list of search terms used.
    """
    from datetime import UTC, datetime

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc: Any = Document()

    # Apply document margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    title = doc.add_heading("Kildeanalyse", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"Procedure: {procedure}")
    run.font.size = Pt(12)
    run.font.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Timestamp
    ts = datetime.now(UTC).replace(microsecond=0).isoformat()
    ts_para = doc.add_paragraph()
    ts_run = ts_para.add_run(f"Genereret: {ts}")
    ts_run.font.size = Pt(10)
    ts_run.font.color.rgb = RGBColor(128, 128, 128)
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # ==========================================================================
    # SECTION 1: READABLE PROSE
    # ==========================================================================

    doc.add_heading("Søgestrategi", level=1)

    # Count sources by type
    source_types: dict[str, int] = {}
    for src in sources:
        kind = src.kind or "unknown"
        source_types[kind] = source_types.get(kind, 0) + 1

    # Prose description of search
    search_para = doc.add_paragraph()
    search_para.add_run(
        f"For at finde evidens til proceduren \"{procedure}\" blev følgende kilder gennemsøgt:\n\n"
    )

    if "danish_guideline" in source_types:
        search_para.add_run(
            f"• Det danske kliniske bibliotek ({source_types['danish_guideline']} guidelines fundet)\n"
        )
    if "pubmed" in source_types:
        search_para.add_run(
            f"• PubMed/MEDLINE ({source_types['pubmed']} artikler fundet)\n"
        )
    if "system_note" in source_types:
        search_para.add_run(
            "• Bemærk: Nogle kilder kunne ikke hentes i denne kørsel\n"
        )

    if search_terms:
        terms_para = doc.add_paragraph()
        terms_para.add_run("Søgetermer anvendt: ").bold = True
        terms_para.add_run(", ".join(search_terms))

    # Evidence hierarchy explanation
    doc.add_heading("Evidenshierarki", level=1)

    hierarchy_para = doc.add_paragraph()
    hierarchy_para.add_run(
        "Kilderne er prioriteret efter evidenshierarkiet:\n\n"
        "1. Danske kliniske guidelines (højeste prioritet) - Disse er udarbejdet af danske "
        "faglige selskaber og regioner specifikt til dansk praksis.\n\n"
        "2. Systematiske reviews og meta-analyser - Giver samlet evidens fra flere studier.\n\n"
        "3. Randomiserede kontrollerede studier (RCT) - Guldstandard for klinisk evidens.\n\n"
        "4. Observationsstudier og case-serier - Lavere evidensniveau, men relevant for sjældne tilstande.\n\n"
        "5. Ekspertudtalelser - Bruges når højere evidens ikke er tilgængelig."
    )

    # Found sources summary
    doc.add_heading("Fundne Kilder", level=1)

    if not sources or (len(sources) == 1 and sources[0].kind == "system_note"):
        no_sources = doc.add_paragraph()
        no_sources.add_run(
            "Ingen eksterne kilder kunne findes til denne procedure. "
            "Dette kan skyldes at søgetermerne ikke matchede tilgængelige dokumenter, "
            "eller at der ikke findes danske guidelines for dette emne."
        )
    else:
        real_sources = [s for s in sources if s.kind != "system_note"]
        for i, src in enumerate(real_sources[:10], 1):  # Limit to 10 in prose
            src_para = doc.add_paragraph()
            src_para.paragraph_format.left_indent = Inches(0.25)

            # Title and type
            src_para.add_run(f"{i}. {src.title or 'Uden titel'}").bold = True
            src_para.add_run(f" ({src.kind})\n")

            # Year if available
            if src.year:
                src_para.add_run(f"   År: {src.year}\n")

            # Evidence level if available
            evidence_level = src.extra.get("evidence_level", "")
            if evidence_level:
                src_para.add_run(f"   Evidensniveau: {evidence_level}\n")

            # Scoring reasoning (first line only)
            reasoning = src.extra.get("scoring_reasoning", [])
            if reasoning and len(reasoning) > 0:
                src_para.add_run(f"   Vurdering: {reasoning[-1]}\n")  # Total score

        if len(real_sources) > 10:
            more_para = doc.add_paragraph()
            more_para.add_run(f"... og {len(real_sources) - 10} yderligere kilder (se teknisk appendix)")

    # ==========================================================================
    # SECTION 2: TECHNICAL APPENDIX
    # ==========================================================================

    doc.add_page_break()
    doc.add_heading("Teknisk Appendix", level=1)

    appendix_intro = doc.add_paragraph()
    appendix_intro.add_run(
        "Dette appendix indeholder detaljerede tekniske data til audit og reproducerbarhed."
    )
    appendix_intro.runs[0].font.italic = True

    # Source scoring table
    doc.add_heading("Kildescore-tabel", level=2)

    real_sources = [s for s in sources if s.kind != "system_note"]
    if real_sources:
        # Create table with headers
        table = doc.add_table(rows=len(real_sources) + 1, cols=5)
        table.style = "Table Grid"

        headers = ["Kilde-ID", "Type", "Samlet Score", "Kvalitet", "Aktualitet"]
        for j, header in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        for i, src in enumerate(real_sources, start=1):
            table.cell(i, 0).text = src.source_id
            table.cell(i, 1).text = src.kind or "unknown"
            table.cell(i, 2).text = f"{src.extra.get('composite_score', 0):.1f}"
            table.cell(i, 3).text = f"{src.extra.get('quality_score', 0):.2f}"
            table.cell(i, 4).text = f"{src.extra.get('recency_score', 0):.2f}"

        doc.add_paragraph()  # Spacer

        # Detailed reasoning for each source
        doc.add_heading("Detaljeret Scoring-begrundelse", level=2)

        for src in real_sources[:20]:  # Limit to 20
            src_heading = doc.add_paragraph()
            src_heading.add_run(f"{src.source_id}: {src.title or 'Uden titel'}").bold = True

            reasoning = src.extra.get("scoring_reasoning", [])
            if reasoning:
                for reason in reasoning:
                    reason_para = doc.add_paragraph()
                    reason_para.paragraph_format.left_indent = Inches(0.25)
                    reason_para.add_run(f"• {reason}")
            else:
                no_reason = doc.add_paragraph()
                no_reason.paragraph_format.left_indent = Inches(0.25)
                no_reason.add_run("Ingen scoring-begrundelse tilgængelig")

            doc.add_paragraph()  # Spacer between sources
    else:
        no_data = doc.add_paragraph()
        no_data.add_run("Ingen kilder at vise i teknisk appendix.")

    # Footer
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = f"Kildeanalyse | run_id={run_id}"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(output_path))


def write_evidence_review_docx(
    *,
    evidence_report: dict[str, Any],
    sources: list[SourceRecord],
    procedure: str,
    run_id: str,
    output_path: Path,
) -> None:
    """Write evidence review documentation as DOCX.

    Hybrid format: Readable prose for clinicians + technical appendix for audit.

    Args:
        evidence_report: Evidence report data (from evidence_report.json).
        sources: List of source records for reference.
        procedure: Name of the procedure being documented.
        run_id: Unique run identifier.
        output_path: Path to save the DOCX file.
    """
    from datetime import UTC, datetime

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc: Any = Document()

    # Apply document margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    title = doc.add_heading("Evidensgennemgang", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"Procedure: {procedure}")
    run.font.size = Pt(12)
    run.font.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Timestamp
    ts = datetime.now(UTC).replace(microsecond=0).isoformat()
    ts_para = doc.add_paragraph()
    ts_run = ts_para.add_run(f"Genereret: {ts}")
    ts_run.font.size = Pt(10)
    ts_run.font.color.rgb = RGBColor(128, 128, 128)
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # ==========================================================================
    # SECTION 1: READABLE PROSE
    # ==========================================================================

    doc.add_heading("Verifikationsproces", level=1)

    process_para = doc.add_paragraph()
    process_para.add_run(
        "Hver påstand i den genererede procedure er blevet systematisk verificeret mod de "
        "tilgængelige kilder. Denne proces sikrer, at procedurens indhold er evidensbaseret "
        "og kan spores tilbage til kliniske kilder.\n\n"
        "Verifikationen klassificerer hver påstand som:\n"
        "• Understøttet - Påstanden har direkte støtte i en eller flere kilder\n"
        "• Delvist understøttet - Dele af påstanden har støtte, men ikke alt\n"
        "• Ikke understøttet - Ingen kilder understøtter påstanden direkte\n"
        "• Modstridende - En kilde modsiger påstanden"
    )

    # Results summary
    doc.add_heading("Resultater", level=1)

    sentences = evidence_report.get("sentences", [])
    sentence_count = evidence_report.get("sentence_count", len(sentences))

    supported_count = sum(1 for s in sentences if s.get("supported", False))
    unsupported_count = sentence_count - supported_count

    # Calculate percentages
    if sentence_count > 0:
        support_pct = (supported_count / sentence_count) * 100
    else:
        support_pct = 0

    results_para = doc.add_paragraph()
    results_para.add_run(f"Antal analyserede påstande: {sentence_count}\n\n")
    results_para.add_run(f"✓ Understøttede påstande: {supported_count} ({support_pct:.0f}%)\n").bold = True
    results_para.add_run(f"✗ Ikke-understøttede påstande: {unsupported_count} ({100-support_pct:.0f}%)\n")

    # Quality assessment
    doc.add_heading("Kvalitetsvurdering", level=1)

    quality_para = doc.add_paragraph()

    if support_pct >= 80:
        quality_para.add_run("Høj evidensdækning: ").bold = True
        quality_para.add_run(
            f"Proceduren har {support_pct:.0f}% understøttede påstande, hvilket indikerer "
            "solid evidensbasering. Proceduren kan anvendes med høj tillid til det faglige indhold."
        )
    elif support_pct >= 50:
        quality_para.add_run("Moderat evidensdækning: ").bold = True
        quality_para.add_run(
            f"Proceduren har {support_pct:.0f}% understøttede påstande. "
            "Visse dele af proceduren bør gennemgås af en klinisk ekspert for at verificere "
            "de ikke-understøttede påstande."
        )
    else:
        quality_para.add_run("Lav evidensdækning: ").bold = True
        quality_para.add_run(
            f"Kun {support_pct:.0f}% af påstandene er understøttet af kilder. "
            "Proceduren kræver omfattende klinisk review før anvendelse. "
            "Overvej at søge efter yderligere kilder eller konsultere fageksperter."
        )

    # Recommendations
    doc.add_heading("Anbefalinger", level=1)

    rec_para = doc.add_paragraph()
    if unsupported_count > 0:
        rec_para.add_run("Baseret på evidensgennemgangen anbefales følgende:\n\n")

        if unsupported_count > sentence_count * 0.5:
            rec_para.add_run("1. Søg efter yderligere danske eller internationale kilder\n")
            rec_para.add_run("2. Konsultér relevant fagligt selskab for ekspertudtalelser\n")
            rec_para.add_run("3. Overvej at simplificere proceduren til kun understøttede elementer\n")
        else:
            rec_para.add_run("1. Gennemgå de ikke-understøttede påstande med klinisk ekspert\n")
            rec_para.add_run("2. Tilføj eksplicitte noter hvor evidens mangler\n")
            rec_para.add_run("3. Opdatér proceduren når nye kilder bliver tilgængelige\n")
    else:
        rec_para.add_run(
            "Alle påstande er understøttet af kilder. Proceduren er klar til klinisk brug "
            "efter standard godkendelsesproces."
        )

    # ==========================================================================
    # SECTION 2: TECHNICAL APPENDIX
    # ==========================================================================

    doc.add_page_break()
    doc.add_heading("Teknisk Appendix", level=1)

    appendix_intro = doc.add_paragraph()
    appendix_intro.add_run(
        "Dette appendix indeholder detaljerede verifikationsdata for hver påstand."
    )
    appendix_intro.runs[0].font.italic = True

    # Unsupported claims table
    unsupported = [s for s in sentences if not s.get("supported", False)]

    if unsupported:
        doc.add_heading("Ikke-understøttede Påstande", level=2)

        for i, claim in enumerate(unsupported[:30], 1):  # Limit to 30
            claim_para = doc.add_paragraph()
            claim_para.paragraph_format.left_indent = Inches(0.25)

            text = claim.get("clean_text", claim.get("text", ""))
            if len(text) > 200:
                text = text[:200] + "..."

            claim_para.add_run(f"{i}. ").bold = True
            claim_para.add_run(text)

            # Show line number if available
            line_no = claim.get("line_no")
            if line_no:
                claim_para.add_run(f" (linje {line_no})")
                claim_para.runs[-1].font.color.rgb = RGBColor(128, 128, 128)

        if len(unsupported) > 30:
            more = doc.add_paragraph()
            more.add_run(f"... og {len(unsupported) - 30} yderligere ikke-understøttede påstande")

    # Supported claims with citations
    supported = [s for s in sentences if s.get("supported", False)]

    if supported:
        doc.add_heading("Understøttede Påstande (med kilder)", level=2)

        for i, claim in enumerate(supported[:20], 1):  # Limit to 20
            claim_para = doc.add_paragraph()
            claim_para.paragraph_format.left_indent = Inches(0.25)

            text = claim.get("clean_text", claim.get("text", ""))
            if len(text) > 150:
                text = text[:150] + "..."

            claim_para.add_run(f"{i}. ").bold = True
            claim_para.add_run(text)

            # Show citations
            citations = claim.get("citations", [])
            if citations:
                claim_para.add_run(f" [{', '.join(citations)}]")
                claim_para.runs[-1].font.color.rgb = RGBColor(0, 100, 0)

        if len(supported) > 20:
            more = doc.add_paragraph()
            more.add_run(f"... og {len(supported) - 20} yderligere understøttede påstande")

    # Summary statistics table
    doc.add_heading("Statistik-oversigt", level=2)

    stats_table = doc.add_table(rows=4, cols=2)
    stats_table.style = "Table Grid"

    stats_data = [
        ("Samlet antal påstande", str(sentence_count)),
        ("Understøttede", f"{supported_count} ({support_pct:.1f}%)"),
        ("Ikke-understøttede", f"{unsupported_count} ({100-support_pct:.1f}%)"),
        ("Antal kilder anvendt", str(len([s for s in sources if s.kind != "system_note"]))),
    ]

    for i, (label, value) in enumerate(stats_data):
        stats_table.cell(i, 0).text = label
        stats_table.cell(i, 1).text = value
        stats_table.cell(i, 0).paragraphs[0].runs[0].bold = True

    # Footer
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = f"Evidensgennemgang | run_id={run_id}"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(output_path))
